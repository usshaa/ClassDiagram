from docx import Document
from docx.shared import Inches

# Create a new document
doc = Document()

# Add a heading
doc.add_heading('Class Diagram', level=1)

# Add an image to the document
doc.add_picture('classdg.png', width=Inches(6))  # Replace 'class_diagram.png' with the actual file name of your generated class diagram image

# Save the document
doc.save('classdg.docx')
